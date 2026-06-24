"""
civitas.ingestion.parsers.docx
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
DOCX parser using python-docx.
Preserves headings and table content.
"""

from __future__ import annotations

import io
import logging

from civitas.ingestion.connectors.base import RawDocumentBlob
from civitas.ingestion.parsers.base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """
    Microsoft Word DOCX parser.

    Extracts:
      · Paragraph text in document order
      · Headings (used as section markers)
      · Table cell content (flattened)
    """

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def parse(self, blob: RawDocumentBlob) -> ParseResult:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = DocxDocument(io.BytesIO(blob.content_bytes))
        parts: list[str] = []
        headings: list[str] = []
        tables_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                headings.append(text)
                parts.append(f"\n## {text}\n")
            else:
                parts.append(text)

        for table in doc.tables:
            tables_count += 1
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                parts.append("\n[TABLE]\n" + "\n".join(rows_text) + "\n[/TABLE]\n")

        return ParseResult(
            text="\n\n".join(parts),
            section_headings=headings,
            tables_count=tables_count,
        )
